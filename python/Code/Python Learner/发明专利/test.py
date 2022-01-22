import re
from docx import Document

re1 = r"\d{1,4}-\w{1,10}"
re2 = r"\w{1,10}"
# 从文件创建文档对象
document = Document('t.docx')
# 显示每段的内容
for p in document.paragraphs:
    res = re.findall(re1, p.text)
    if res:
        break


def replace_text(doc, old_text, new_text, model_text):
    # 遍历每个段落
    exist = 0;
    for p in doc.paragraphs:
        # 如果要搜索的内容在该段落
        if model_text not in p.text:
            if old_text in p.text:
                exist = 1;
                # 注意！runs 会根据样式分隔内容，确保被替换内容的样式一致
                p.text = p.text.replace(old_text, new_text)
                p.text = p.text.replace("所述所述", "所述")
    if exist == 0:
        print("WARNING   "+old_text+"是未使用的专有名词")


for s in res:
    res2 = re.findall(re2, s)
    replace_text(document, res2[1], "所述"+res2[1]+res2[0], res2[0]+"-"+res2[1])

document.save('result.docx')
print("")
input('输入任意按键结束程序运行')
